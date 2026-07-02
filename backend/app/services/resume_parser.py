"""
Production-grade Resume Parser.
Extracts structured data from PDF and DOCX files.
"""

import re
import json
from pathlib import Path
from typing import Optional, Dict, Any, List
from loguru import logger

import fitz  # PyMuPDF
import pdfplumber
from docx import Document


# ── Skill Keywords (extensible) ───────────────────────────────────────────────
PROGRAMMING_SKILLS = {
    "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust",
    "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "sql", "bash",
    "shell", "perl", "lua", "dart", "julia",
}

FRAMEWORK_SKILLS = {
    "react", "angular", "vue", "nextjs", "nuxtjs", "svelte", "django", "flask",
    "fastapi", "spring", "express", "nestjs", "laravel", "rails", "gin",
    "fastify", "graphql", "rest", "grpc",
}

ML_SKILLS = {
    "machine learning", "deep learning", "nlp", "computer vision", "pytorch",
    "tensorflow", "keras", "scikit-learn", "sklearn", "pandas", "numpy",
    "huggingface", "langchain", "rag", "llm", "gpt", "bert", "transformers",
    "opencv", "xgboost", "lightgbm", "catboost",
}

CLOUD_SKILLS = {
    "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "k8s",
    "terraform", "ansible", "ci/cd", "github actions", "jenkins", "gitlab ci",
    "heroku", "vercel", "netlify", "cloudflare",
}

DATABASE_SKILLS = {
    "postgresql", "mysql", "mongodb", "redis", "elasticsearch", "sqlite",
    "cassandra", "dynamodb", "neo4j", "bigquery", "snowflake", "pinecone",
    "chromadb", "weaviate", "firebase",
}

ALL_SKILLS = (
    PROGRAMMING_SKILLS
    | FRAMEWORK_SKILLS
    | ML_SKILLS
    | CLOUD_SKILLS
    | DATABASE_SKILLS
)


class ResumeParser:
    """Parses PDF and DOCX resumes into structured JSON."""

    def parse(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Main entry point. Parse a resume file and return structured data.

        Args:
            file_path: Absolute path to the file.
            file_type: 'pdf' or 'docx'

        Returns:
            Structured dictionary with all extracted fields.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        if file_type.lower() == "pdf":
            raw_text = self._extract_text_pdf(str(path))
        elif file_type.lower() == "docx":
            raw_text = self._extract_text_docx(str(path))
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        logger.bind(chars=len(raw_text).info("Resume text extracted"), file_type=file_type)
        return self._parse_text(raw_text)

    # ── Text Extraction ────────────────────────────────────────────────────────

    def _extract_text_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF first, fallback to pdfplumber."""
        try:
            text_parts = []
            with fitz.open(file_path) as doc:
                for page in doc:
                    text_parts.append(page.get_text("text"))
            text = "\n".join(text_parts).strip()
            if len(text) > 100:
                return text
        except Exception as e:
            logger.bind(error=str(e).warning("PyMuPDF extraction failed, trying pdfplumber"))

        # Fallback to pdfplumber
        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return "\n".join(text_parts).strip()
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {e}") from e

    def _extract_text_docx(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            return "\n".join(paragraphs).strip()
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {e}") from e

    # ── Parsing Logic ──────────────────────────────────────────────────────────

    def _parse_text(self, text: str) -> Dict[str, Any]:
        """Parse raw resume text into structured data."""
        lines = [l.strip() for l in text.split("\n") if l.strip()]

        return {
            "raw_text": text,
            "full_name": self._extract_name(lines),
            "email": self._extract_email(text),
            "phone": self._extract_phone(text),
            "location": self._extract_location(text),
            "linkedin_url": self._extract_url(text, "linkedin"),
            "github_url": self._extract_url(text, "github"),
            "skills": self._extract_skills(text),
            "experience": self._extract_experience(text, lines),
            "education": self._extract_education(text, lines),
            "projects": self._extract_projects(text, lines),
            "certifications": self._extract_certifications(text, lines),
            "languages": self._extract_languages(text),
            "total_years_experience": self._estimate_years_experience(text),
            "seniority_level": self._estimate_seniority(text),
        }

    def _extract_name(self, lines: List[str]) -> Optional[str]:
        """Extract full name from first few lines."""
        name_pattern = re.compile(r"^[A-Z][a-z]+(?: [A-Z][a-z]+){1,3}$")
        for line in lines[:5]:
            if name_pattern.match(line.strip()):
                return line.strip()
        return None

    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address."""
        pattern = re.compile(r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b")
        match = pattern.search(text)
        return match.group(0).lower() if match else None

    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number."""
        patterns = [
            r"\+?[\d\s\-\(\)]{10,16}",
            r"\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b",
        ]
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                phone = re.sub(r"[^\d+]", "", match.group(0))
                if 10 <= len(phone) <= 15:
                    return match.group(0).strip()
        return None

    def _extract_location(self, text: str) -> Optional[str]:
        """Extract location from text."""
        patterns = [
            r"(?:Location|Address|Based in)[:\s]+([^\n,]+(?:,\s*[^\n]+)?)",
            r"\b([A-Z][a-z]+(?:,\s*[A-Z]{2,})?(?:,\s*[A-Z][a-z]+)?)\b",
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return None

    def _extract_url(self, text: str, platform: str) -> Optional[str]:
        """Extract a specific platform URL."""
        pattern = re.compile(
            rf"https?://(?:www\.)?{re.escape(platform)}\.com/[^\s\"'<>]+",
            re.IGNORECASE,
        )
        match = pattern.search(text)
        return match.group(0) if match else None

    def _extract_skills(self, text: str) -> Dict[str, List[str]]:
        """Extract skills by category."""
        text_lower = text.lower()
        return {
            "programming": [s for s in PROGRAMMING_SKILLS if re.search(rf"\b{re.escape(s)}\b", text_lower)],
            "frameworks": [s for s in FRAMEWORK_SKILLS if re.search(rf"\b{re.escape(s)}\b", text_lower)],
            "ml_ai": [s for s in ML_SKILLS if re.search(rf"\b{re.escape(s)}\b", text_lower)],
            "cloud_devops": [s for s in CLOUD_SKILLS if re.search(rf"\b{re.escape(s)}\b", text_lower)],
            "databases": [s for s in DATABASE_SKILLS if re.search(rf"\b{re.escape(s)}\b", text_lower)],
        }

    def _extract_experience(self, text: str, lines: List[str]) -> List[Dict[str, Any]]:
        """Extract work experience sections."""
        experiences = []
        exp_pattern = re.compile(
            r"(?:EXPERIENCE|WORK EXPERIENCE|EMPLOYMENT|PROFESSIONAL EXPERIENCE)",
            re.IGNORECASE,
        )
        date_pattern = re.compile(
            r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[.\s,]+\d{4}"
            r"(?:\s*[-–—]\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[.\s,]+\d{4}|"
            r"\s*[-–—]\s*(?:Present|Current|Now))?",
            re.IGNORECASE,
        )

        in_section = False
        current_entry: Dict[str, Any] = {}
        bullet_points: List[str] = []

        for line in lines:
            if exp_pattern.search(line):
                in_section = True
                continue
            if in_section and re.match(r"(?:EDUCATION|PROJECTS|SKILLS|CERTIFICATIONS)", line, re.IGNORECASE):
                if current_entry:
                    current_entry["responsibilities"] = bullet_points
                    experiences.append(current_entry)
                break

            if in_section:
                date_match = date_pattern.search(line)
                if date_match:
                    if current_entry:
                        current_entry["responsibilities"] = bullet_points
                        experiences.append(current_entry)
                        bullet_points = []
                    current_entry = {
                        "title": line.replace(date_match.group(0), "").strip(),
                        "duration": date_match.group(0).strip(),
                        "company": "",
                    }
                elif current_entry and line.startswith(("•", "-", "*", "·")):
                    bullet_points.append(line.lstrip("•-*· ").strip())

        if current_entry:
            current_entry["responsibilities"] = bullet_points
            experiences.append(current_entry)

        return experiences

    def _extract_education(self, text: str, lines: List[str]) -> List[Dict[str, Any]]:
        """Extract education entries."""
        education = []
        edu_pattern = re.compile(r"(?:EDUCATION|ACADEMIC)", re.IGNORECASE)
        degree_pattern = re.compile(
            r"\b(?:B\.?Tech|B\.?E|B\.?Sc|B\.?A|M\.?Tech|M\.?Sc|M\.?S|M\.?B\.?A|Ph\.?D|Bachelor|Master|Doctor)\b",
            re.IGNORECASE,
        )

        in_section = False
        for line in lines:
            if edu_pattern.search(line):
                in_section = True
                continue
            if in_section and re.match(
                r"(?:EXPERIENCE|PROJECTS|SKILLS|CERTIFICATIONS)", line, re.IGNORECASE
            ):
                break
            if in_section and degree_pattern.search(line):
                year_match = re.search(r"\b(20\d{2}|19\d{2})\b", line)
                education.append({
                    "degree": degree_pattern.search(line).group(0),
                    "institution": line,
                    "year": year_match.group(0) if year_match else None,
                })

        return education

    def _extract_projects(self, text: str, lines: List[str]) -> List[Dict[str, Any]]:
        """Extract project entries."""
        projects = []
        proj_pattern = re.compile(r"(?:PROJECTS|PROJECT EXPERIENCE)", re.IGNORECASE)
        tech_pattern = re.compile(r"(?:Tech Stack|Technologies|Built with|Stack)[:\s]+([^\n]+)", re.IGNORECASE)

        in_section = False
        current_proj: Dict[str, Any] = {}
        desc_lines: List[str] = []

        for line in lines:
            if proj_pattern.search(line):
                in_section = True
                continue
            if in_section and re.match(
                r"(?:EXPERIENCE|EDUCATION|SKILLS|CERTIFICATIONS)", line, re.IGNORECASE
            ):
                if current_proj:
                    current_proj["description"] = " ".join(desc_lines)
                    projects.append(current_proj)
                break

            if in_section:
                tech_match = tech_pattern.search(line)
                if tech_match:
                    if current_proj:
                        current_proj["description"] = " ".join(desc_lines)
                        projects.append(current_proj)
                        desc_lines = []
                    current_proj = {
                        "name": line.split(":")[0].strip() if ":" in line else line.strip(),
                        "tech_stack": [t.strip() for t in tech_match.group(1).split(",")],
                    }
                elif current_proj:
                    desc_lines.append(line.lstrip("•-*· ").strip())

        if current_proj:
            current_proj["description"] = " ".join(desc_lines)
            projects.append(current_proj)

        return projects

    def _extract_certifications(self, text: str, lines: List[str]) -> List[str]:
        """Extract certifications."""
        certs = []
        cert_pattern = re.compile(
            r"\b(?:AWS|Azure|GCP|Google|Oracle|Cisco|PMP|Scrum|CISSP|CPA|CFA|"
            r"TensorFlow|PyTorch|Databricks|Snowflake|Kubernetes|CKAD|CKA)\b",
            re.IGNORECASE,
        )
        in_section = False
        for line in lines:
            if re.match(r"(?:CERTIFICATIONS|CERTIFICATES|LICENSES)", line, re.IGNORECASE):
                in_section = True
                continue
            if in_section and re.match(
                r"(?:EXPERIENCE|EDUCATION|PROJECTS|SKILLS)", line, re.IGNORECASE
            ):
                break
            if in_section and cert_pattern.search(line):
                certs.append(line.strip())

        return certs

    def _extract_languages(self, text: str) -> List[str]:
        """Extract human languages."""
        known_languages = {
            "english", "hindi", "french", "german", "spanish", "japanese",
            "chinese", "arabic", "portuguese", "russian", "korean", "italian",
        }
        text_lower = text.lower()
        return [lang.capitalize() for lang in known_languages if lang in text_lower]

    def _estimate_years_experience(self, text: str) -> float:
        """Estimate total years of professional experience from dates."""
        date_pattern = re.compile(
            r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[.\s,]+(\d{4})",
            re.IGNORECASE,
        )
        years = [int(m.group(1)) for m in date_pattern.finditer(text)]
        if len(years) >= 2:
            return float(max(years) - min(years))
        return 0.0

    def _estimate_seniority(self, text: str) -> str:
        """Estimate seniority level from years of experience."""
        years = self._estimate_years_experience(text)
        if years >= 8:
            return "senior"
        elif years >= 3:
            return "mid"
        elif years >= 1:
            return "junior"
        return "entry"
